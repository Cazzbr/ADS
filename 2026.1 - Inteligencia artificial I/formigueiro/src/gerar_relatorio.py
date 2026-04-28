from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    p._element.get_or_add_pPr().append(shading)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def gerar():
    doc = Document()

    # ---------------------------------------------------------------
    # Capa
    # ---------------------------------------------------------------
    title = doc.add_heading(
        "Trabalho Prático: Forrageamento e Comunicação em Enxames\n"
        "(Simulação de Formigueiro)",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Disciplina: Inteligência Artificial I  |  ADS 2026.1\n"
        "Autor: Luciano Magri\n"
        "Instituição: IFRS — Campus Farroupilha\n"
        "Data: Abril de 2026"
    )
    r.font.size = Pt(11)
    r.italic = True

    doc.add_page_break()

    # ---------------------------------------------------------------
    # 1. Introdução
    # ---------------------------------------------------------------
    add_heading(doc, "1. Introdução")
    add_paragraph(
        doc,
        "Este relatório documenta a adaptação de um sistema multiagente baseado em "
        "Q-Learning para agendamento de horários para uma simulação de forrageamento de "
        "formigueiro. A nova implementação modela formigas como agentes autônomos e "
        "assíncronos (Python asyncio) que buscam alimento em uma grade 10×10, "
        "aplicando conceitos de busca distribuída, feromônio negativo e protocolo de "
        "comunicação por passagem de mensagens.",
    )

    # ---------------------------------------------------------------
    # 2. Mudanças Realizadas
    # ---------------------------------------------------------------
    add_heading(doc, "2. Mudanças Realizadas na Adaptação")

    add_heading(doc, "2.1 Remoção do Q-Learning", level=2)
    add_paragraph(
        doc,
        "O código original utilizava uma tabela Q (q_table), política epsilon-greedy e "
        "cálculo de recompensas para que agentes negociassem horários sem conflito. "
        "Toda essa lógica foi removida pois o problema do formigueiro não envolve "
        "seleção de slots em um domínio fixo, mas sim navegação espacial em um grafo.",
    )
    add_paragraph(doc, "Elementos removidos:")
    add_bullet(doc, "Atributos: q_table, alpha, epsilon, appointment, agent_view, prioridades.")
    add_bullet(doc, "Métodos: escolher_horario(), _conflitos_com_mesmo_horario(), _ha_superior_em_conflito(), _horarios_bloqueados_por_superiores().")
    add_bullet(doc, "Toda lógica de reward shaping e decaimento de epsilon.")

    add_heading(doc, "2.2 Nova Classe: Formiga", level=2)
    add_paragraph(
        doc,
        "A classe QLearningAgent foi substituída pela classe Formiga "
        "(base_formiga.py). Os atributos centrais passaram a ser:",
    )
    add_bullet(doc, "position: posição atual na grade (tupla x, y).")
    add_bullet(doc, "route: histórico de coordenadas da tentativa atual.")
    add_bullet(doc, "known_food_path: rota de sucesso recebida de outra formiga (None enquanto desconhecida).")
    add_bullet(doc, "inbox: asyncio.Queue() — mantido do código original para receber mensagens assíncronas.")

    add_heading(doc, "2.3 Nova Classe: Ambiente", level=2)
    add_paragraph(
        doc,
        "Foi introduzida a classe Ambiente (formiga.py) para centralizar o estado "
        "compartilhado do mundo:",
    )
    add_bullet(doc, "grade NxN com vizinhança 4-conectada.")
    add_bullet(doc, "pos_alimento: posição do alimento, oculta das formigas.")
    add_bullet(doc, "alimento_coletado / capacidade_maxima: contador de coletas (esgota em 4).")
    add_bullet(doc, "caminhos_invalidos: lista global de rotas fracassadas (feromônio negativo).")

    add_heading(doc, "2.4 Reformulação das Mensagens", level=2)
    add_paragraph(
        doc,
        "O protocolo de mensagens foi completamente reformulado. No sistema original, "
        "mensagens carregavam slot_escolhido para resolução de conflitos de agendamento. "
        "No novo sistema existem apenas dois tipos de mensagem:",
    )
    add_bullet(doc, "{'tipo': 'sucesso', 'de': nome, 'rota': [...]} — broadcast ao encontrar alimento.")
    add_bullet(doc, "{'tipo': 'fracasso', 'de': nome, 'rota': [...]} — broadcast ao atingir limite de passos.")

    add_heading(doc, "2.5 Reformulação do Orquestrador", level=2)
    add_paragraph(
        doc,
        "O executor_qlearning.py verificava convergência de slots únicos entre agentes. "
        "O executor_formigueiro.py monitora exclusivamente o contador de coletas "
        "(ambiente.alimento_coletado) e cancela as tarefas quando atingir 4, "
        "sem interferir nas decisões de rota das formigas.",
    )

    # ---------------------------------------------------------------
    # 3. Lógica Aplicada
    # ---------------------------------------------------------------
    add_heading(doc, "3. Lógica Aplicada — Tarefas de Implementação")

    add_heading(doc, "Tarefa A: Modelagem do Ambiente e Movimentação", level=2)
    add_paragraph(
        doc,
        "O ambiente é modelado como uma grade 10×10 com coordenadas inteiras. "
        "O formigueiro fica em (0,0) e o alimento é posicionado aleatoriamente "
        "a distância Manhattan entre 3 e 12 passos, tornando a busca desafiante "
        "mas viável dentro do limite de 20 passos. A corrotina _forragear() executa "
        "o loop de movimentação:",
    )
    add_code_block(
        doc,
        "for passo in range(MAX_PASSOS):\n"
        "    await self._processar_inbox()\n"
        "    next_pos = random.choice(self._vizinhos_validos())\n"
        "    self.position = next_pos\n"
        "    self.route.append(next_pos)\n"
        "    if self.ambiente.tem_comida(self.position):\n"
        "        await self._retornar_sucesso(); return\n"
        "    await asyncio.sleep(0.02)",
    )
    add_paragraph(
        doc,
        "O método _vizinhos_validos() retorna as coordenadas 4-adjacentes dentro "
        "dos limites da grade, excluindo as que aparecem como próximo passo em "
        "caminhos inválidos conhecidos.",
    )

    add_heading(doc, "Tarefa B: Retorno e Feromônio Negativo", level=2)
    add_paragraph(
        doc,
        "Quando a formiga esgota os 20 passos sem encontrar alimento, invoca "
        "_retornar_fracasso():",
    )
    add_bullet(doc, "Registra a rota completa em ambiente.caminhos_invalidos (feromônio negativo global).")
    add_bullet(doc, "Faz broadcast da mensagem 'fracasso' para todas as outras formigas.")
    add_bullet(
        doc,
        "As formigas receptoras adicionam a rota à lista global e passam a evitar "
        "o próximo passo daquele caminho ao escolher vizinhos — comportamento análogo "
        "ao Nogood do algoritmo Asynchronous Backtracking (ABT).",
    )

    add_heading(doc, "Tarefa C: Comunicação de Sucesso (Broadcast)", level=2)
    add_paragraph(
        doc,
        "Ao encontrar o alimento, a formiga invoca _retornar_sucesso():",
    )
    add_bullet(doc, "Incrementa ambiente.alimento_coletado.")
    add_bullet(doc, "Faz broadcast da mensagem 'sucesso' com a rota completa (formigueiro → alimento).")
    add_bullet(
        doc,
        "As formigas receptoras armazenam a rota em known_food_path e passam a "
        "segui-la diretamente nas próximas tentativas, sem exploração aleatória — "
        "inspirado no algoritmo LRTA(n)*, onde agentes compartilham as melhores "
        "rotas encontradas para acelerar a convergência coletiva.",
    )

    add_heading(doc, "Tarefa D: Orquestrador e Condição de Parada", level=2)
    add_paragraph(
        doc,
        "O executor_formigueiro.py cria uma corrotina por formiga e monitora "
        "passivamente o contador de coletas em um loop com asyncio.sleep(0.1). "
        "Quando ambiente.alimento_coletado atinge 4 (ou o timeout de 120 s é "
        "atingido), todas as tarefas são canceladas via tarefa.cancel() e "
        "asyncio.gather(*tarefas, return_exceptions=True) aguarda o encerramento "
        "gracioso. Em seguida é impresso o relatório de execução.",
    )

    # ---------------------------------------------------------------
    # 4. Análise Crítica
    # ---------------------------------------------------------------
    add_heading(doc, "4. Análise Crítica")

    add_heading(doc, "4.1 Exploração vs. Explotação", level=2)
    add_paragraph(
        doc,
        "Enquanto nenhuma formiga conhece a localização do alimento, o enxame opera "
        "em modo de exploração pura: cada agente executa um passeio aleatório "
        "independente, cobrindo regiões diferentes da grade de forma emergente. "
        "No momento em que uma formiga bem-sucedida transmite a rota exata, o "
        "sistema faz uma transição coletiva para explotação: as demais formigas "
        "absorvem o conhecimento via inbox e passam a seguir o caminho conhecido "
        "diretamente, sem desvio.",
    )
    add_paragraph(
        doc,
        "Comparando com agentes que aprendem de forma isolada: um agente isolado "
        "precisaria redescobrir o alimento a cada tentativa (ou manter sua própria "
        "memória), com custo proporcional ao número de tentativas individuais. "
        "No sistema coletivo, uma única descoberta beneficia imediatamente todo "
        "o enxame — o custo de exploração é amortizado entre os agentes. Isso é "
        "análogo ao reforço positivo de feromônios nas formigas reais: a trilha "
        "bem-sucedida é reforçada (aqui, pela rota compartilhada), convergindo o "
        "comportamento do grupo para a solução ótima sem que cada indivíduo precise "
        "aprender do zero.",
    )

    add_heading(doc, "4.2 Impacto do Reforço Negativo", level=2)
    add_paragraph(
        doc,
        "O feromônio negativo (lista caminhos_invalidos) atua como poda do espaço "
        "de busca: a cada tentativa fracassada, um subconjunto de arestas do grafo "
        "é marcado como improdutivo e evitado pelas demais formigas. Nos logs da "
        "simulação é possível observar que, após as primeiras tentativas fracassadas, "
        "as formigas seguintes desviam das direções já exploradas e exploram regiões "
        "novas da grade, reduzindo redundância.",
    )
    add_paragraph(
        doc,
        "Contudo, há uma limitação importante de escala. Em um mapa muito grande "
        "(ex.: grade 50×50), o número de caminhos possíveis de até 20 passos cresce "
        "exponencialmente (na ordem de 4^20 ≈ 10^12). Bloquear apenas as rotas já "
        "percorridas não garante cobertura suficiente do espaço de busca, e o alimento "
        "pode nunca ser encontrado dentro do limite de energia. Estratégias "
        "complementares seriam necessárias: aumentar o limite de passos dinamicamente, "
        "usar busca em leque (múltiplas colônias com pontos de partida distribuídos), "
        "ou combinar feromônio negativo com feromônio positivo de exploração "
        "(reforço por novidade) para direcionar ativamente as formigas a regiões "
        "não visitadas.",
    )

    add_heading(doc, "4.3 Comunicação Direta vs. Estigmergia", level=2)
    add_paragraph(
        doc,
        "Na implementação atual, formigas comunicam-se diretamente via passagem de "
        "mensagens (asyncio.Queue): a formiga bem-sucedida envia a rota para a "
        "inbox de cada colega. Na natureza, formigas alteram o ambiente depositando "
        "feromônios no chão — comunicação indireta conhecida como estigmergia.",
    )
    add_paragraph(
        doc,
        "Para implementar estigmergia mantendo a arquitetura atual, seriam necessárias "
        "apenas duas modificações:",
    )
    add_bullet(
        doc,
        "Na classe Ambiente: ",
        bold_prefix="Mapa de feromônios — ",
    )
    add_paragraph(
        doc,
        "        adicionar mapa_feromonios: dict[tuple, float] inicializado em 0.0 "
        "para todas as células. Formigas bem-sucedidas depositam feromônio ao longo "
        "de sua rota (incremento proporcional à qualidade do caminho). O orquestrador "
        "aplica evaporação periódica multiplicando todos os valores por um fator "
        "ρ < 1 a cada iteração.",
    )
    add_bullet(
        doc,
        "Na classe Formiga: ",
        bold_prefix="Seleção probabilística — ",
    )
    add_paragraph(
        doc,
        "        substituir random.choice(candidatos) em _vizinhos_validos() por "
        "uma seleção por roleta (roulette wheel) baseada nos valores de feromônio "
        "de cada vizinho candidato (maior feromônio = maior probabilidade de escolha). "
        "Remover completamente os métodos _broadcast() e _processar_inbox(), pois "
        "a comunicação passa a ser totalmente mediada pelo ambiente.",
    )
    add_paragraph(
        doc,
        "Com essas mudanças, nenhum agente conhece diretamente a decisão de outro; "
        "a coordenação emerge apenas da modificação e leitura do estado do ambiente — "
        "princípio fundamental da estigmergia.",
    )

    # ---------------------------------------------------------------
    # 5. Conclusão
    # ---------------------------------------------------------------
    add_heading(doc, "5. Conclusão")
    add_paragraph(
        doc,
        "A adaptação demonstrou como a arquitetura assíncrona por passagem de "
        "mensagens, originalmente usada para resolução de conflitos de agendamento "
        "com Q-Learning, pode ser reutilizada para modelar comportamentos emergentes "
        "de enxame. Os mecanismos de feromônio negativo (análogo ao Nogood do ABT) "
        "e de broadcast de sucesso (análogo ao LRTA(n)*) implementam, de forma "
        "simplificada, os dois pilares do Ant Colony Optimization: evitar soluções "
        "ruins e reforçar as boas. O sistema converge satisfatoriamente em grades "
        "pequenas, mas escalabilidade em ambientes maiores requereria a migração "
        "para estigmergia com mapa de feromônios, eliminando a necessidade de "
        "comunicação direta entre agentes.",
    )

    output_path = "relatorio_formigueiro.docx"
    doc.save(output_path)
    print(f"Relatório gerado: {output_path}")


if __name__ == "__main__":
    gerar()
