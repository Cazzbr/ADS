from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    p._element.get_or_add_pPr().append(shading)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def gerar():
    doc = Document()

    # ---------------------------------------------------------------
    # Capa
    # ---------------------------------------------------------------
    title = doc.add_heading(
        "Forrageamento e Comunicação em Enxames\nSimulação de Formigueiro",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Trabalho Prático de Inteligência Artificial I\n"
        "Disciplina: Sistemas Multiagente Assíncrono\n"
        "Autor: Luciano Magri\n"
        "IFRS — Campus Farroupilha\n"
        "Abril de 2026"
    )
    r.font.size = Pt(11)
    r.italic = True

    doc.add_page_break()

    # ---------------------------------------------------------------
    # Resumo Executivo
    # ---------------------------------------------------------------
    add_heading(doc, "Resumo Executivo")
    add_paragraph(
        doc,
        "Este trabalho implementa uma simulação multiagente assíncrona de forrageamento "
        "em formigueiro utilizando Python asyncio. Cinco formigas autônomas exploram "
        "uma grade 10×10 em busca de alimento oculto. Uma vez descoberto, a formiga "
        "bem-sucedida comunica a rota exata às demais, que imediatamente convergem para "
        "a exploração coletiva do recurso. O sistema implementa dois mecanismos de "
        "aprendizado distribuído: (i) reforço negativo via marcação de caminhos "
        "inválidos, análogo a nogoods em algoritmos de busca com restrições; e "
        "(ii) reforço positivo via broadcast de rotas bem-sucedidas, análogo ao "
        "compartilhamento de políticas em LRTA(n)*. A simulação converge com sucesso, "
        "coletando 4 unidades de alimento em menos de 30 segundos de simulação.",
    )

    # ---------------------------------------------------------------
    # 1. Objetivo
    # ---------------------------------------------------------------
    add_heading(doc, "1. Objetivo e Contexto")
    add_paragraph(
        doc,
        "Modelar o comportamento emergente de um enxame de formigas buscando alimento "
        "em um ambiente desconhecido. Cada formiga é um agente autônomo que:",
    )
    add_bullet(doc, "Executa caminhada aleatória com limite de energia (20 passos).")
    add_bullet(doc, "Marca caminhos fracassados como restrições globais (feromônio negativo).")
    add_bullet(doc, "Compartilha rotas de sucesso com todo o enxame ao encontrar alimento.")
    add_bullet(doc, "Adapta seu comportamento conforme aprende com os peers.")

    add_paragraph(
        doc,
        "O objetivo é demonstrar como comunicação síncrona e estado global compartilhado "
        "podem coordenar agentes para resolver coletivamente um problema que seria "
        "inviável de forma isolada.",
    )

    # ---------------------------------------------------------------
    # 2. Arquitetura
    # ---------------------------------------------------------------
    add_heading(doc, "2. Arquitetura do Sistema")

    add_heading(doc, "2.1 Estrutura de Arquivos", level=2)
    add_bullet(doc, "base_formiga.py: Classe Formiga com lógica de movimentação e comunicação.")
    add_bullet(doc, "formiga.py: Classe Ambiente e instanciação de agentes.")
    add_bullet(doc, "executor_formigueiro.py: Orquestrador que monitora o estado global.")

    add_heading(doc, "2.2 Componentes Principais", level=2)
    add_bullet(
        doc,
        "Classe Ambiente: ",
        bold_prefix="",
    )
    add_paragraph(
        doc,
        "Mantém a grade NxN, posição do alimento (oculta dos agentes), contador de coletas, "
        "e lista global de caminhos inválidos (feromônio negativo). Oferece interfaces para "
        "verificação de alimento e obtenção de vizinhos válidos.",
    )
    add_bullet(
        doc,
        "Classe Formiga: ",
        bold_prefix="",
    )
    add_paragraph(
        doc,
        "Agente autônomo com inbox (asyncio.Queue) para recebimento de mensagens. "
        "Mantém histórico da rota atual, rota conhecida (broadcast recebido), e posição. "
        "A corrotina agir() executa tentativas infinitas de forrageamento.",
    )
    add_bullet(
        doc,
        "Orquestrador (main): ",
        bold_prefix="",
    )
    add_paragraph(
        doc,
        "Cria corrotinas para cada formiga e monitora o contador de coletas de forma "
        "passiva. Não dita rotas; apenas cancela todas as tarefas quando alimento_coletado == 4.",
    )

    # ---------------------------------------------------------------
    # 3. Implementação das Tarefas
    # ---------------------------------------------------------------
    add_heading(doc, "3. Implementação das Tarefas (A, B, C, D)")

    add_heading(doc, "Tarefa A: Movimentação Assíncrona", level=2)
    add_paragraph(
        doc,
        "A corrotina _forragear() executa a busca em um loop de até 20 passos:",
    )
    add_code_block(
        doc,
        "for passo in range(MAX_PASSOS):\n"
        "    await self._processar_inbox()  # verifica mensagens\n"
        "    if self.known_food_path:  # segue rota conhecida se disponível\n"
        "        next_pos = self.known_food_path[passo + 1]\n"
        "    else:\n"
        "        next_pos = random.choice(self._vizinhos_validos())\n"
        "    self.position = next_pos\n"
        "    self.route.append(next_pos)\n"
        "    if self.ambiente.tem_comida(self.position):\n"
        "        await self._retornar_sucesso(); return\n"
        "    await asyncio.sleep(0.02)",
    )
    add_paragraph(
        doc,
        "O método _vizinhos_validos() retorna vizinhos 4-adjacentes dentro da grade, "
        "evitando aqueles que aparecem em caminhos inválidos conhecidos (Tarefa B).",
    )

    add_heading(doc, "Tarefa B: Reforço Negativo (Feromônio)", level=2)
    add_paragraph(
        doc,
        "Ao atingir 20 passos sem alimento, _retornar_fracasso() executa:",
    )
    add_bullet(
        doc,
        "Registra a rota completa em ambiente.caminhos_invalidos (estado global).",
    )
    add_bullet(
        doc,
        "Envia broadcast do tipo 'fracasso' com a rota para todas as outras formigas.",
    )
    add_bullet(
        doc,
        "Formigas receptoras adicionam a rota à lista global e passam a evitar "
        "o próximo passo daquele caminho ao escolher vizinhos.",
    )
    add_paragraph(
        doc,
        "Este mecanismo é análogo ao Nogood do algoritmo Asynchronous Backtracking (ABT), "
        "onde um agente comunica aos demais uma atribuição (rota) que não leva a solução válida.",
    )

    add_heading(doc, "Tarefa C: Comunicação de Sucesso", level=2)
    add_paragraph(
        doc,
        "Ao encontrar alimento, _retornar_sucesso() executa:",
    )
    add_bullet(
        doc,
        "A própria formiga memoriza a rota (self.known_food_path = rota).",
    )
    add_bullet(
        doc,
        "Incrementa o contador global (ambiente.alimento_coletado).",
    )
    add_bullet(
        doc,
        "Envia broadcast do tipo 'sucesso' com a rota exata para todas as demais formigas.",
    )
    add_bullet(
        doc,
        "Formigas receptoras armazenam a rota e passam a seguir diretamente nas próximas "
        "tentativas, sem exploração aleatória.",
    )
    add_paragraph(
        doc,
        "Este padrão é inspirado no algoritmo LRTA(n)*, onde agentes exploram um ambiente "
        "e compartilham as melhores rotas encontradas para acelerar convergência coletiva.",
    )

    add_heading(doc, "Tarefa D: Orquestrador", level=2)
    add_paragraph(
        doc,
        "O executor cria uma corrotina por formiga e monitora passivamente:",
    )
    add_code_block(
        doc,
        "while ambiente.alimento_coletado < ambiente.capacidade_maxima:\n"
        "    if timeout_exceeded:\n"
        "        break\n"
        "    await asyncio.sleep(0.1)",
    )
    add_paragraph(
        doc,
        "Quando alimento_coletado atinge 4 (ou timeout), todas as tarefas são canceladas "
        "graciosamente com tarefa.cancel().",
    )

    # ---------------------------------------------------------------
    # 4. Análise Crítica
    # ---------------------------------------------------------------
    add_heading(doc, "4. Análise Crítica e Conceitual")

    add_heading(doc, "4.1 Exploração vs. Explotação", level=2)
    add_paragraph(
        doc,
        "O sistema exibe transição clara entre dois modos operacionais:",
    )
    add_paragraph(
        doc,
        "Fase de Exploração: Antes de qualquer formiga encontrar alimento, o enxame "
        "opera em modo de exploração pura — cada agente executa passeio aleatório independente, "
        "cobrindo regiões diferentes de forma emergente.",
    )
    add_paragraph(
        doc,
        "Fase de Explotação: Uma vez que uma formiga bem-sucedida compartilha a rota, "
        "o sistema faz transição coletiva para explotação — todas as demais formigas absorvem "
        "o conhecimento via inbox e passam a seguir o caminho conhecido diretamente, sem desvio.",
    )
    add_paragraph(
        doc,
        "Comparação com aprendizado isolado: um agente isolado precisaria redescobrir o "
        "alimento a cada tentativa individual. No sistema coletivo, o custo de exploração "
        "é amortizado — uma única descoberta beneficia imediatamente todo o enxame. Isso é "
        "análogo ao reforço positivo de feromônios em formigas reais: a trilha bem-sucedida "
        "é reforçada (aqui, pela rota compartilhada), convergindo comportamento do grupo para "
        "a solução ótima sem que cada indivíduo aprenda isoladamente.",
    )

    add_heading(doc, "4.2 Impacto do Reforço Negativo", level=2)
    add_paragraph(
        doc,
        "O feromônio negativo (lista caminhos_invalidos) atua como poda do espaço de busca. "
        "A cada tentativa fracassada, um subconjunto de arestas do grafo é marcado como "
        "improdutivo. Nos logs, é observável que após tentativas iniciais fracassadas, "
        "formigas seguintes desviam de direções já exploradas e exploram regiões novas, "
        "reduzindo redundância significativamente.",
    )
    add_paragraph(
        doc,
        "Limitação de escala: Em mapas muito grandes (ex.: 50×50), o número de caminhos "
        "possíveis de até 20 passos cresce exponencialmente (~4^20 ≈ 10^12). Bloquear apenas "
        "rotas percorridas não garante cobertura suficiente, e alimento pode nunca ser encontrado. "
        "Estratégias complementares seriam necessárias: aumentar limite dinâmico, usar busca "
        "em leque com múltiplas colônias, ou combinar feromônio negativo com feromônio positivo "
        "de exploração para direcionar ativamente formigas a regiões não visitadas.",
    )

    add_heading(doc, "4.3 Comunicação Direta vs. Estigmergia", level=2)
    add_paragraph(
        doc,
        "Implementação atual usa passagem direta de mensagens (asyncio.Queue). Na natureza, "
        "formigas alteram o ambiente depositando feromônios — comunicação indireta (estigmergia).",
    )
    add_paragraph(
        doc,
        "Para implementar estigmergia, seriam necessárias duas mudanças principais:",
    )
    add_bullet(
        doc,
        "Mapa de feromônios em Ambiente: dict[tuple, float] inicializado em 0.0. "
        "Formigas bem-sucedidas depositam feromônio ao longo da rota (incremento proporcional). "
        "Orquestrador aplica evaporação periódica multiplicando valores por fator ρ < 1.",
    )
    add_bullet(
        doc,
        "Seleção probabilística em Formiga: Substituir random.choice() por roleta "
        "(roulette wheel) baseada em valores de feromônio dos vizinhos (maior feromônio = "
        "maior probabilidade). Remover completamente _broadcast() e _processar_inbox().",
    )
    add_paragraph(
        doc,
        "Com essas mudanças, nenhum agente conhece diretamente decisão de outro; coordenação "
        "emerge apenas de modificação e leitura do estado do ambiente — princípio fundamental "
        "de estigmergia.",
    )

    # ---------------------------------------------------------------
    # 5. Resultados
    # ---------------------------------------------------------------
    add_heading(doc, "5. Resultados e Observações")
    add_paragraph(
        doc,
        "Simulações executadas confirmam comportamento esperado:",
    )
    add_bullet(
        doc,
        "Convergência: Alimento é descoberto e coletado por 4 formigas em menos de 30 segundos.",
    )
    add_bullet(
        doc,
        "Transição de exploração para explotação: Visível nos logs — primeiras tentativas "
        "mostram modo 'aleatório', após broadcast há mudança abrupta para 'rota conhecida'.",
    )
    add_bullet(
        doc,
        "Eficiência coletiva: Uma descoberta beneficia imediatamente todo o enxame; "
        "não há redescoberta individual.",
    )
    add_bullet(
        doc,
        "Feromônio negativo: Bloqueios de caminhos aparecem nos logs; ants demonstram "
        "seletividade após registro de rotas fracassadas.",
    )

    # ---------------------------------------------------------------
    # 6. Conclusão
    # ---------------------------------------------------------------
    add_heading(doc, "6. Conclusão")
    add_paragraph(
        doc,
        "A simulação de formigueiro implementa, de forma didática mas funcional, os dois "
        "pilares do Ant Colony Optimization: evitar soluções ruins (feromônio negativo, "
        "análogo a Nogood) e reforçar as boas (broadcast de sucesso, análogo a LRTA(n)*). "
        "O sistema demonstra como mecanismos simples de comunicação assíncrona podem gerar "
        "comportamento emergente complexo e eficiente em solução de problemas distribuídos.",
    )
    add_paragraph(
        doc,
        "Enquanto a comunicação direta por mensagens funciona bem em escala pequena, "
        "a migração para estigmergia (comunicação indireta via modificação de ambiente) "
        "seria essencial para escalabilidade em cenários reais, eliminando necessidade "
        "de conhecimento global e permitindo operação totalmente distribuída.",
    )

    output_path = "Relatorio_Formigueiro_Final.docx"
    doc.save(output_path)
    print(f"✓ Relatório gerado com sucesso: {output_path}")


if __name__ == "__main__":
    gerar()
