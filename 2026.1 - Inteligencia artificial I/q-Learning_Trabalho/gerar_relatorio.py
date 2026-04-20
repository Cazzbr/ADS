from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Estilos globais ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bold_body(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def code_block(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['No Spacing']
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    return p

# ── Título ───────────────────────────────────────────────────────────────────
title = doc.add_heading('Relatório Comparativo: Algoritmo de Propagação vs. Q-Learning', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Disciplina: ').bold = True
meta.add_run('Inteligência Artificial I     ')
meta.add_run('Data: ').bold = True
meta.add_run('20/04/2026     ')
meta.add_run('Autor: ').bold = True
meta.add_run('Luciano Magri')
doc.add_paragraph()

# ── 1. Introdução ────────────────────────────────────────────────────────────
heading('1. Introdução', 1)
body(
    'Este relatório apresenta uma análise comparativa entre dois algoritmos de resolução '
    'distribuída de conflitos em sistemas multiagente:'
)
bullet('Algoritmo de Propagação (pasta cores/): abordagem reativa baseada em difusão de estado e resolução local de conflitos.')
bullet('Q-Learning (pasta raiz): abordagem de aprendizado por reforço onde agentes aprendem uma política de escolha de horários ao longo do tempo.')
body(
    'Ambos os sistemas utilizam comunicação assíncrona com asyncio e operam sobre um domínio '
    'de 5 valores com 5 agentes em topologia totalmente conectada (grafo completo).'
)

# ── 2. Descrição dos Algoritmos ──────────────────────────────────────────────
heading('2. Descrição dos Algoritmos', 1)

heading('2.1 Algoritmo de Propagação', 2)
body(
    'O algoritmo de propagação implementa um modelo clássico de Coloração de Grafos Distribuída. '
    'Cada agente escolhe aleatoriamente uma cor inicial e, ao detectar conflito com um vizinho, '
    'aplica uma regra determinística de desempate: o agente com o maior nome lexicográfico cede '
    'e escolhe uma nova cor aleatória dentre as não bloqueadas.'
)
bold_body('Fluxo de execução:', '')
for step in [
    'Cada agente é instanciado em seu próprio arquivo (agente_prop_A0.py … agente_prop_A4.py).',
    'O executor conecta as filas de mensagens e dispara as corrotinas.',
    'Cada agente difunde seu estado para vizinhos com nome maior (self.name < n).',
    'Ao receber mensagem, atualiza sua visão local (views) e chama resolve_conflict().',
    'Se houve mudança de cor, re-difunde o novo estado.',
    'O sistema para após um stop_event acionado por timeout fixo de 5 segundos.',
]:
    bullet(step)
bold_body('Formato da mensagem:', '')
for line in ['{"kind": "state", "sender": "A0", "color": "Azul"}']:
    code_block(line)

heading('2.2 Q-Learning', 2)
body(
    'O algoritmo de Q-Learning aplica Aprendizado por Reforço ao problema de alocação de horários. '
    'Cada agente mantém uma tabela Q que associa cada horário do domínio a um valor aprendido, '
    'e utiliza a política epsilon-greedy para equilibrar exploração e explotação.'
)
bold_body('Fluxo de execução:', '')
for step in [
    'Os agentes são instanciados centralmente em agente_qlearning.py com prioridades explícitas.',
    'O executor conecta as filas e dispara a propagação inicial.',
    'Ao receber mensagem, o agente calcula o conflito, determina o próximo horário, computa a recompensa composta, atualiza a tabela Q e propaga se houver mudança.',
    'O epsilon decai a cada atualização (ε = max(0.01, ε × 0.99)), reduzindo gradualmente a exploração.',
    'O sistema para ao detectar ausência de conflitos ou por timeout de 5 segundos.',
]:
    bullet(step)
bold_body('Formato da mensagem:', '')
for line in [
    '{',
    '    "from": "P0",',
    '    "slot_escolhido": "10:00",',
    '    "prioridade": 1,',
    '    "timestamp": 1234567890.0',
    '}',
]:
    code_block(line)

# ── 3. Comparação Técnica ────────────────────────────────────────────────────
heading('3. Comparação Técnica', 1)

headers = ['Critério', 'Propagação', 'Q-Learning']
rows = [
    ('Paradigma',             'Reativo / CSP distribuído',           'Aprendizado por Reforço'),
    ('Domínio',               'Cores (Branco, Preto, Azul…)',        'Horários (09:00 – 15:00)'),
    ('Aprendizado',           'Nenhum — regra determinística',        'Tabela Q atualizada a cada iteração'),
    ('Política de decisão',   'Desempate lexicográfico por nome',     'Epsilon-greedy com decaimento'),
    ('Prioridade',            'Implícita (nome lexicográfico)',        'Explícita (número inteiro por agente)'),
    ('Recompensa',            'Inexistente',                          '+10 único, −20 conflito, −2 mudança, +5 prioridade'),
    ('Comunicação',           'Difusão seletiva (nome menor envia)',  'Difusão completa (todos os vizinhos)'),
    ('Latência simulada',     'Sim (0.03–0.18 s por mensagem)',       'Não'),
    ('Critério de parada',    'Timeout fixo via stop_event',          'Convergência detectada + timeout'),
    ('Exploração',            'Escolha aleatória entre cores livres', 'Epsilon-greedy com ε decaindo até 0.01'),
    ('Métricas coletadas',    'Nenhuma',                              'Conflitos/iteração, epsilon médio'),
]

table = doc.add_table(rows=1 + len(rows), cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), '1F497D')
    hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

for i, (criterio, prop, ql) in enumerate(rows):
    row_cells = table.rows[i + 1].cells
    row_cells[0].text = criterio
    row_cells[1].text = prop
    row_cells[2].text = ql
    if i % 2 == 0:
        for cell in row_cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), 'EEF3F8')
            cell._tc.get_or_add_tcPr().append(shading)

doc.add_paragraph()

# ── 4. Estrutura de Comunicação ──────────────────────────────────────────────
heading('4. Estrutura de Comunicação', 1)

heading('4.1 Propagação', 2)
body(
    'A comunicação é assimétrica: apenas o agente com nome lexicograficamente menor inicia a '
    'difusão (self.name < n), evitando mensagens duplicadas no par. Isso reduz o tráfego na '
    'rede à metade em topologia completa.'
)

heading('4.2 Q-Learning', 2)
body(
    'A comunicação é simétrica e completa: ao mudar de horário, o agente notifica todos os vizinhos. '
    'A mensagem inclui metadados extras (prioridade, timestamp) que permitem rastreabilidade e '
    'poderiam ser usados para lógica de desempate baseada em chegada.'
)

# ── 5. Resolução de Conflitos ────────────────────────────────────────────────
heading('5. Resolução de Conflitos', 1)

heading('5.1 Propagação', 2)
body(
    'A resolução é puramente local e determinística. Ao detectar conflito, o agente verifica se '
    'seu nome é o menor dentre os conflitantes. Se for, mantém a cor; caso contrário, escolhe '
    'aleatoriamente uma cor não bloqueada pelos vizinhos. Não há memória de decisões anteriores.'
)

heading('5.2 Q-Learning', 2)
body(
    'A resolução é orientada por aprendizado e prioridade numérica. O agente conflitante com um '
    'vizinho de maior prioridade (número menor) é obrigado a ceder. O valor Q do horário atual '
    'é atualizado com a recompensa composta, fazendo com que o agente aprenda a preferir horários '
    'que historicamente geram menos conflito.'
)

# ── 6. Vantagens e Desvantagens ──────────────────────────────────────────────
heading('6. Vantagens e Desvantagens', 1)

heading('6.1 Propagação', 2)
p = doc.add_paragraph()
p.add_run('Vantagens:').bold = True
for v in [
    'Simples de implementar e entender.',
    'Convergência rápida em grafos esparsos.',
    'Sem hiperparâmetros para ajustar.',
    'Arquitetura descentralizada verdadeira (cada agente em seu próprio arquivo).',
]:
    bullet(v)
p = doc.add_paragraph()
p.add_run('Desvantagens:').bold = True
for d in [
    'Sem aprendizado — não melhora com o tempo.',
    'Convergência baseada em timeout, não em detecção real.',
    'Não coleta métricas de desempenho.',
    'A regra lexicográfica pode gerar oscilações em grafos densos.',
]:
    bullet(d)

heading('6.2 Q-Learning', 2)
p = doc.add_paragraph()
p.add_run('Vantagens:').bold = True
for v in [
    'Agentes aprendem a preferir horários com menor histórico de conflitos.',
    'Sistema de recompensa expressivo e configurável.',
    'Convergência detectada ativamente.',
    'Coleta de métricas para análise gráfica (conflitos por iteração, epsilon médio).',
    'Prioridades explícitas permitem hierarquia controlada.',
]:
    bullet(v)
p = doc.add_paragraph()
p.add_run('Desvantagens:').bold = True
for d in [
    'Mais complexo: requer ajuste de α, ε e taxa de decaimento.',
    'Q-Learning sem transição de estado é tecnicamente um algoritmo bandit, não Q-Learning pleno.',
    'Convergência não garantida dentro do timeout em espaços com alta exploração.',
]:
    bullet(d)

# ── 7. Conclusão ─────────────────────────────────────────────────────────────
heading('7. Conclusão', 1)
body(
    'Os dois algoritmos abordam o mesmo problema — alocação sem conflito em ambiente distribuído — '
    'por caminhos opostos. O algoritmo de propagação é eficiente, previsível e adequado para '
    'problemas estáticos onde a topologia e o domínio são conhecidos. O Q-Learning é mais adequado '
    'para ambientes dinâmicos onde os agentes precisam adaptar seu comportamento com base em '
    'experiência acumulada, pagando o custo de maior complexidade e necessidade de ajuste de '
    'hiperparâmetros.'
)
body(
    'Para o contexto acadêmico apresentado (5 agentes, domínio fixo de 5 valores), ambos convergem '
    'de forma satisfatória. Em cenários maiores ou com domínios variáveis, o Q-Learning tenderia a '
    'apresentar vantagem adaptativa, enquanto o algoritmo de propagação manteria vantagem em '
    'velocidade de convergência inicial.'
)

doc.add_paragraph()
footer = doc.add_paragraph('Relatório gerado automaticamente com base na análise do código-fonte dos algoritmos.')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
footer.runs[0].font.size = Pt(9)

out = 'relatorio_comparativo.docx'
doc.save(out)
print(f'Salvo: {out}')
